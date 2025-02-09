from typing import Optional

from PySide6.QtWidgets import QMainWindow
from PySide6.QtCore import Qt, QSettings

from autoscore.mainui import Ui_MainWindowAutoScore
from autoscore.safemath import safe_eval_math
from autoscore.version import get_version
from loguru import logger as log
from pathlib import Path
import re
import pandas as pd
import statistics
from collections import namedtuple
from docx import Document


PaperStats = namedtuple(
    "PaperStats", ["File", "N", "Scores", "Mean", "MeanAdj", "Dropped", "PPS", "PPSAdj", "Additions", "Note"]
)


class MainWin(QMainWindow):
    def __init__(self):
        super(MainWin, self).__init__()
        self.data = []

        # Precompile regex patterns
        self.score_regex = re.compile(r"\[([\d\.\-\+\*\/\(\)]+)\]", re.IGNORECASE)
        self.addition_regex = re.compile(r"{([-+]*[\d\.]+)\|([^|}]+)}", re.IGNORECASE)
        self.addition_ranged_regex = re.compile(r"{([-+]*[\d\.]+)\|([\w ]+)\|([\d\.]+)\|([\d\.]+)}", re.IGNORECASE)

        self.ui = Ui_MainWindowAutoScore()
        self.ui.setupUi(self)

        self.load_settings()
        self.setup_ui_connections()

        self.ui.labelTitle.setText(
            f"""
        <html>
            <head/>
                <body>
                    <p align="center">
                        <span style="font-size:16pt; font-weight:600;">
                            AutoScore Files v{get_version()}
                        </span>
                    </p>
                    <p align="center">
                        <span style="font-size:12pt;">Travis L. Seymour, PhD</span>
                    </p>
                    <hr/>
                    <p>Drag/Drop Files or Folders here (.txt and .docx only)</p>
                </body>
        </html>
        """
        )

        self.ui.pushButtonQuit.clicked.connect(self.on_quit)
        self.ui.pushButtonClearList.clicked.connect(self.on_clear_list)
        self.ui.pushButtonProcess.clicked.connect(self.process_list)
        self.ui.lineEditAdjustments.textChanged.connect(self.check_adjustment_entry)
        self.ui.listWidgetFileList.setActiveCallback(self.set_ui_buttons)

    def load_settings(self):
        settings = QSettings()
        self.ui.checkBoxAppendToFile.setChecked(settings.value("checkBoxAppendToFile", False, type=bool))
        self.ui.checkBoxEndAfterProcess.setChecked(settings.value("checkBoxEndAfterProcess", False, type=bool))
        self.ui.lineEditAdjustments.setText(settings.value("lineEditAdjustments", "", type=str))
        self.ui.spinboxScoresToDrop.setValue(settings.value("spinboxScoresToDrop", 0, type=int))

    def save_settings(self):
        settings = QSettings()
        settings.setValue("checkBoxAppendToFile", self.ui.checkBoxAppendToFile.isChecked())
        settings.setValue("checkBoxEndAfterProcess", self.ui.checkBoxEndAfterProcess.isChecked())
        settings.setValue("lineEditAdjustments", self.ui.lineEditAdjustments.text())
        settings.setValue("spinboxScoresToDrop", self.ui.spinboxScoresToDrop.value())

    def setup_ui_connections(self):
        self.ui.checkBoxAppendToFile.toggled.connect(self.save_settings)
        self.ui.checkBoxEndAfterProcess.toggled.connect(self.save_settings)
        self.ui.lineEditAdjustments.textChanged.connect(self.save_settings)
        self.ui.spinboxScoresToDrop.valueChanged.connect(self.save_settings)

    def set_ui_buttons(self, active: bool = True):
        self.ui.pushButtonProcess.setEnabled(active)
        self.ui.pushButtonClearList.setEnabled(active)

    def on_clear_list(self):
        self.data.clear()
        self.ui.listWidgetFileList.clear()
        self.set_ui_buttons(False)

    def on_quit(self):
        self.close()

    def check_adjustment_entry(self, text):
        self.ui.labelAdjustments.setEnabled(
            bool(self.addition_regex.findall(text) or self.addition_ranged_regex.findall(text))
        )

    @staticmethod
    def get_file_contents(file_path: Path) -> str:
        if file_path.suffix.lower() in (".txt", ".prs"):
            return file_path.read_text()
        elif file_path.suffix.lower() == ".docx":
            from docx2txt import process

            return process(str(file_path))
        return ""

    @staticmethod
    def rawscore(txtscore):
        try:
            return safe_eval_math(txtscore.strip().strip("[]"))
        except (ValueError, SyntaxError):
            log.warning(f"Invalid score format: {txtscore}")
            return None

    def analyze(self, file_name: str, contents: str) -> PaperStats:
        try:
            if not contents:
                raise ValueError(f"File {Path(file_name).name} appears to be empty!")

            scoreList = [self.rawscore(match.group(1)) for match in self.score_regex.finditer(contents)]
            scoreList = [x for x in scoreList if x is not None]

            if not scoreList:
                raise ValueError(f"No score markers found in {Path(file_name).name} (e.g., [98] or [70+25])")

            scoreList_adj = scoreList.copy()
            Dropped = self.ui.spinboxScoresToDrop.value()
            for _ in range(Dropped):
                if len(scoreList_adj) > 1:
                    scoreList_adj.remove(min(scoreList_adj))

            N = len(scoreList)
            Scores = "|".join(map(str, scoreList))
            Mean = statistics.mean(scoreList) if scoreList else float("nan")
            MeanAdj = statistics.mean(scoreList_adj) if scoreList_adj else float("nan")
            PPS = 100 / N if N else float("nan")
            PPSAdj = 100 / len(scoreList_adj) if scoreList_adj else float("nan")
            Note = ""

            if self.ui.labelAdjustments.isEnabled():
                _contents = f"{self.ui.lineEditAdjustments.text().strip()}\n\n{contents}"
            else:
                _contents = contents

            ranged_additions = self.addition_ranged_regex.findall(_contents)
            additions = self.addition_regex.findall(_contents)

            Additions = [(float(adj), msg, 0.0, 200.0) for adj, msg in additions] + [
                (float(adj), msg, float(lo), float(hi)) for adj, msg, lo, hi in ranged_additions
            ]
            Additions = [adj for adj in Additions if adj[2] <= MeanAdj <= adj[3]]
            for adjustment in Additions:
                MeanAdj += adjustment[0]

            return PaperStats(file_name, N, Scores, Mean, MeanAdj, Dropped, PPS, PPSAdj, Additions, Note)

        except ValueError as e:
            self.statusBar().showMessage(str(e))
            return PaperStats(
                file_name, 0, "", float("nan"), float("nan"), float("nan"), float("nan"), float("nan"), tuple(), str(e)
            )

    def save_data(self, start_path: Path):
        output_file = start_path / "0_grading_results.xlsx"
        df = pd.DataFrame(self.data)
        df["Mean"] = df["Mean"].round(0)
        df["MeanAdj"] = df["MeanAdj"].round(0)
        df.to_excel(output_file, index=False, engine="openpyxl")

    def append_paper_stats(self, file_path: Path, paper_stats: PaperStats) -> bool:
        additions_text = (
            (
                "\n        Manual Grade Adjustments:\n"
                + "\n".join(
                    [
                        f"                {'+' if adj > 0.0 else ''}{adj}: {msg}"
                        for adj, msg, _, _ in paper_stats.Additions
                    ]
                )
            )
            if paper_stats.Additions
            else ""
        )

        adj_note = f"Dropped lowest {paper_stats.Dropped} scores." if paper_stats.Dropped else "None."
        results = f"""
        Source File: {paper_stats.File}
        Number of Scores: {paper_stats.N}
        Scores: {paper_stats.Scores}
        Score Adjustments: {adj_note} 
        Other Adjustments: {additions_text}
        Grade: {paper_stats.MeanAdj:.0f}%
        """

        try:
            doc = Document(str(file_path))
            doc.add_page_break()
            p0 = doc.add_paragraph()
            p0.add_run("PAPER GRADE INFORMATION").bold = True
            p0.add_run(f"\n({file_path.stem.split('_')[0]})")
            doc.add_paragraph(results)

            out_path = file_path.parent / "updated_with_grade_info"
            out_path.mkdir(exist_ok=True)
            doc.save(str(out_path / file_path.name))

            return True
        except Exception as e:
            log.error(f"Unable to add paper stats to {file_path.name}: {e}")
            return False

    def process_list(self):
        self.data.clear()
        self.statusBar().clearMessage()
        first_path: Optional[Path] = None
        for item in self.ui.listWidgetFileList.findItems("*", Qt.MatchFlag.MatchWildcard):
            file_path = Path(item.text())
            if first_path is None:
                first_path = Path(file_path).absolute()

            if file_path.name == "0_grading_results.xlsx":
                continue
            if not file_path.is_file():
                log.error(f"Could not access file: {file_path}")
                continue

            self.statusBar().showMessage(f"Processing {file_path.name}")
            contents = self.get_file_contents(file_path)
            if not contents:
                log.error(f"Error reading file: {file_path}")
                continue

            result = self.analyze(file_path.name, contents)
            self.data.append(result)

            if self.ui.checkBoxAppendToFile.isChecked():
                self.append_paper_stats(file_path, result)

        if self.data and first_path is not None:
            self.save_data(start_path=file_path.parent)  # Path(self.data[0].File).parent)
            self.statusBar().showMessage('Successfully saved results to "grading_results.xlsx"!')

        if self.ui.checkBoxEndAfterProcess.isChecked():
            self.close()
        else:
            self.on_clear_list()
